#!/usr/bin/env python3
"""Render manuscript_nsclc.md to a plain-black Word .docx.

Pandoc (bundled via pypandoc) handles structure, tables, embedded figures, and
the numbered reference list. A python-docx post-pass forces every style and run
to black so there are no colored Word heading defaults ("basic word format black").

Run from repo root:  python scripts/nsclc/build_manuscript_docx.py
"""
from pathlib import Path
import pypandoc
from docx import Document
from docx.shared import RGBColor

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "docs" / "paper1_nsclc" / "manuscript_nsclc.md"
OUT = ROOT / "docs" / "paper1_nsclc" / "Sociodemographic_Bias_NSCLC_manuscript.docx"
BLACK = RGBColor(0, 0, 0)


def main():
    # 1) pandoc: markdown -> docx. Run with cwd=ROOT so figures/... paths resolve.
    pypandoc.convert_file(
        str(SRC),
        "docx",
        outputfile=str(OUT),
        extra_args=["--from=markdown", f"--resource-path={ROOT}"],
    )

    # 2) force plain black everywhere (styles + runs)
    doc = Document(str(OUT))
    for style in doc.styles:
        font = getattr(style, "font", None)
        if font is not None:
            try:
                font.color.rgb = BLACK
            except Exception:
                pass
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.color.rgb = BLACK
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = BLACK
    doc.save(str(OUT))
    print(f"Wrote {OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
