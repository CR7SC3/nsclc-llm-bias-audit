#!/usr/bin/env python3
"""Build NSCLC_methods_and_figures.docx from methods_and_figures.md.

A minimal, dependency-light renderer (python-docx only). Handles the small
markdown subset this handout uses:
  # ...   -> Title
  ## ...  -> Heading 1
  ### ... -> Heading 2
  blank-line-separated paragraphs -> Normal, with inline **bold** runs.

Run from the repo root:  python scripts/nsclc/build_methods_figures_docx.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "docs" / "paper1_nsclc" / "methods_and_figures.md"
OUT = ROOT / "docs" / "paper1_nsclc" / "NSCLC_methods_and_figures.docx"


def add_runs(paragraph, text):
    """Split text on ** and emit alternating normal/bold runs."""
    for i, seg in enumerate(text.split("**")):
        if seg == "":
            continue
        run = paragraph.add_run(seg)
        run.bold = i % 2 == 1


def blocks(md_text):
    """Yield (kind, text). kind in {title, h1, h2, para}. Paragraphs join
    consecutive non-blank, non-heading lines."""
    buf = []
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if line.startswith("### "):
            if buf:
                yield ("para", " ".join(buf)); buf = []
            yield ("h2", line[4:].strip())
        elif line.startswith("## "):
            if buf:
                yield ("para", " ".join(buf)); buf = []
            yield ("h1", line[3:].strip())
        elif line.startswith("# "):
            if buf:
                yield ("para", " ".join(buf)); buf = []
            yield ("title", line[2:].strip())
        elif line.strip() == "":
            if buf:
                yield ("para", " ".join(buf)); buf = []
        else:
            buf.append(line.strip())
    if buf:
        yield ("para", " ".join(buf))


def main():
    doc = Document()
    # Base body font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for kind, text in blocks(SRC.read_text(encoding="utf-8")):
        if kind == "title":
            doc.add_paragraph(text, style="Title")
        elif kind == "h1":
            doc.add_paragraph(text, style="Heading 1")
        elif kind == "h2":
            doc.add_paragraph(text, style="Heading 2")
        else:
            p = doc.add_paragraph(style="Normal")
            add_runs(p, text)

    doc.save(OUT)
    print(f"Wrote {OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
